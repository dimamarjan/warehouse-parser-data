import json
from docx import Document
from locationVoskresen import locationAddres


def build_file(loaded_file):
    data = []
    table1 = loaded_file.tables[0]
    for i, row in enumerate(table1.rows):
        for text in row.cells:
            data.append(text.text)
    for item in data[0:4:]:
        data.remove(item)
    return data


def split_positions(data_list):
    split_list = []
    temp = []
    for k in data_list:
        temp.append(k)
        if len(temp) == 4:
            split_list.append(temp)
            temp = []
    return split_list


def compare_lists(list1, list2=None):
    if list2 is None:
        list2 = []
    dataList = []

    try:
        location = []
        for position in list1:
            location = position[1].split()
            dataList.append({
                "id": position[0],
                "cargo": bool(0),
                "changeQ": bool(0),
                "changedQTo": int(0),
                "name": position[1],
                "quantity": position[2],
                "location": locationAddres(location[len(location) - 1]),
                "uhtishka": bool(0),
                "uhtQuant": int(0)
            })

        if len(list2) != 0:
            for position in list2:
                location = position[1].split()
                dataList.append({
                    "id": str((int(position[0]) + len(list1))),
                    "cargo": bool(0),
                    "changeQ": bool(0),
                    "changedQTo": int(0),
                    "name": position[1],
                    "quantity": position[2],
                    "location": locationAddres(location[len(location) - 1]),
                    "uhtishka": bool(0),
                    "uhtQuant": int(0)
                })
        buildedData = {"positions": dataList}
        return buildedData
    except IndexError as err:
        dataList = 0
        print("comparing error: ", err)
        buildedData = {"positions": dataList}
        return buildedData


if __name__ == "__main__":
    split_list1 = []
    split_list2 = []
    try:
        document1 = Document('./docx/tempFile-t.docx')
        input_data1 = build_file(document1)
        split_list1 = split_positions(input_data1)
    except BaseException as error:
        print("file-1: ", error)
        input("Файл tempFile-t.docx не найден...")
        exit()
    try:
        document2 = Document("./docx/tempFile-v.docx")
        input_data2 = build_file(document2)
        split_list2 = split_positions(input_data2)
    except BaseException as error:
        print("file-2: ", error)
    if len(split_list2) != 0:
        positions_compared = compare_lists(split_list1, split_list2)
    else:
        positions_compared = compare_lists(split_list1)

    if positions_compared != 0:
        with open("./wh-data-serv/db.json", "w+", encoding="utf-8") as f:
            json.dump(positions_compared, f, ensure_ascii=False)
    else:
        pass

